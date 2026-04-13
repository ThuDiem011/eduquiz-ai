from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_ultimate_report():
    doc = Document()
    
    # --- Style Config ---
    styles = doc.styles
    for level in range(1, 4):
        style = styles[f'Heading {level}']
        style.font.name = 'Times New Roman'
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800

    # --- COVER PAGE ---
    doc.add_paragraph("\n" * 2)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO ĐỒ ÁN CHUYÊN NGÀNH")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'

    prj_p = doc.add_paragraph()
    prj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prj_run = prj_p.add_run("\nTên đề tài:\nEDUQUIZ AI - NỀN TẢNG ÔN LUYỆN TRẮC NGHIỆM THÔNG MINH\nHỖ TRỢ BỞI TRÍ TUỆ NHÂN TẠO (GENERATIVE AI)")
    prj_run.font.size = Pt(26)
    prj_run.font.bold = True
    prj_run.font.name = 'Times New Roman'
    prj_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo 600

    doc.add_paragraph("\n" * 8)
    
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    auth_run = auth_p.add_run("Lĩnh vực: Khoa học máy tính / Trí tuệ nhân tạo\nCông nghệ: Next.js 16 & Gemini AI\n" + f"Ngày hoàn thành: {datetime.date.today().strftime('%d/%m/%Y')}")
    auth_run.font.size = Pt(13)
    auth_run.font.name = 'Times New Roman'

    doc.add_page_break()

    # --- CHƯƠNG 1 ---
    doc.add_heading('CHƯƠNG 1: TỔNG QUAN DỰ ÁN', level=1)
    doc.add_paragraph(
        "Dự án EduQuiz AI ra đời trong bối cảnh cuộc cách mạng công nghiệp 4.0 đang tác động mạnh mẽ đến giáo dục. "
        "Việc ôn tập và tự đánh giá năng lực là nhu cầu thiết yếu của người học, tuy nhiên việc tìm kiếm hoặc tự soạn "
        "đề thi chất lượng tốn rất nhiều thời gian. EduQuiz AI giải quyết bài toán này bằng cách kết hợp sức mạnh "
        "của Generative AI để tạo ra hàng ngàn câu hỏi chuẩn xác, đa dạng và cá nhân hóa lộ trình học tập cho từng sinh viên."
    )

    # --- CHƯƠNG 2 ---
    doc.add_heading('CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG & CẤU TRÚC THƯ MỤC', level=1)
    doc.add_paragraph(
        "Hệ thống được xây dựng theo mô hình Monolith hiện đại với Next.js App Router, giúp tối ưu hóa hiệu năng "
        "và khả năng SEO. Cấu trúc thư mục được phân chia khoa học:"
    )
    doc.add_paragraph(
        "• /src/app: Chứa các Route và các trang giao diện (Pages/Layouts).\n"
        "• /src/api: Cổng kết nối dữ liệu Serverless API xử lý các tác vụ Backend.\n"
        "• /src/components: Tập hợp các thành phần giao diện tái sử dụng (Atomic Design).\n"
        "• /src/lib: Chứa các Service lõi như AI Generator, Analytics Engine và Database Client.\n"
        "• /prisma: Quản lý lược đồ dữ liệu và các tập lệnh di cư (Migrations)."
    )

    # --- CHƯƠNG 3 ---
    doc.add_heading('CHƯƠNG 3: CÔNG NGHỆ & FRAMEWORK SỬ DỤNG', level=1)
    doc.add_paragraph("Hệ thống tận dụng các công nghệ tiên tiến nhất hiện nay:")
    
    techs = [
        ("Framework", "Next.js 16 (Turbopack) - Tăng tốc độ render và dev mode."),
        ("Language", "TypeScript 5 - Đảm bảo an toàn kiểu dữ liệu, giảm thiểu lỗi runtime."),
        ("Styling", "Tailwind CSS - Hệ thống utility-first giúp giao diện nhất quán, hiện đại."),
        ("Database", "Turso DB (LibSQL/SQLite) - Cơ sở dữ liệu phân tán tốc độ cao."),
        ("ORM", "Prisma - Trình quản lý database mạnh mẽ nhất cho Node.js."),
        ("AI Model", "Google Gemini 1.5 Flash - Mô hình ngôn ngữ lớn thế hệ mới với token window rộng.")
    ]
    for name, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # --- CHƯƠNG 4 ---
    doc.add_heading('CHƯƠNG 4: THIẾT KẾ CƠ SỞ DỮ LIỆU (PRISMA SCHEMA)', level=1)
    doc.add_paragraph(
        "Cơ sở dữ liệu được thiết kế chặt chẽ với các quan hệ thực thể quan trọng:\n"
        "- Subject (Môn học) → Chapter (Chương) → Lesson (Bài học).\n"
        "- Question (Câu hỏi) liên kết với nhiều Choice (Lựa chọn).\n"
        "- Exam (Đề thi) tập hợp nhiều ExamQuestion (Bảng trung gian).\n"
        "- StudentExamAttempt (Lịch sử làm bài) ghi lại điểm số và quá trình làm bài của sinh viên."
    )

    # --- CHƯƠNG 5 ---
    doc.add_heading('CHƯƠNG 5: CƠ CHẾ HOẠT ĐỘNG CỦA AI', level=1)
    doc.add_paragraph(
        "Luồng xử lý thông minh của AI Generator:\n"
        "1. Trích xuất ngữ cảnh từ dữ liệu môn học/chương trong hệ thống.\n"
        "2. Áp dụng Prompt Engineering: Định nghĩa vai trò (chuyên gia khảo thí), cấp độ khó và định dạng JSON.\n"
        "3. Cơ chế Context Awareness: AI đọc danh sách các câu hỏi cũ để đảm bảo không tạo trùng lặp.\n"
        "4. Phân tích phản hồi: Hệ thống tự động trích xuất chuỗi JSON từ phản hồi của AI và chuẩn hóa dữ liệu."
    )

    # --- CHƯƠNG 6 ---
    doc.add_heading('CHƯƠNG 6: MODULE PHÂN TÍCH NĂNG LỰC (ANALYTICS)', level=1)
    doc.add_paragraph(
        "Đây là trái tim của hệ thống giúp học sinh tiến bộ. Thuật toán phân tích dựa trên:\n"
        "- Mastery Score: Tính điểm thành thạo theo trọng số độ khó (EASY: 1, MEDIUM: 1.5, HARD: 2).\n"
        "- Analysis Result: Tự động đưa ra nhận xét bằng tiếng Việt và gợi ý nội dung cần tập trung ôn tập."
    )

    # --- CHƯƠNG 7 ---
    doc.add_heading('CHƯƠNG 7: MODULE LEADERBOARD & VINH DANH', level=1)
    doc.add_paragraph(
        "Hệ thống xếp hạng sử dụng tính năng 'Mastery Point' để vinh danh những người dùng tích cực nhất. "
        "Giao diện Leaderboard được thiết kế cao cấp với hiệu ứng bục vinh danh (Podium) cho Top 3, "
        "tạo động lực thi đua học tập."
    )

    # --- CHƯƠNG 8 ---
    doc.add_heading('CHƯƠNG 8: HƯỚNG DẪN CÀI ĐẶT (MÔI TRƯỜNG TRẮNG)', level=1)
    doc.add_paragraph("Để cài đặt hệ thống trên một máy tính hoàn toàn mới, hãy thực hiện các bước sau:")
    doc.add_paragraph(
        "Bước 1: Cài đặt Node.js (phiên bản v18 trở lên) từ nodejs.org.\n"
        "Bước 2: Cài đặt Git (git-scm.com) để quản lý mã nguồn.\n"
        "Bước 3: Tải mã nguồn dự án về máy.\n"
        "Bước 4: Mở Terminal tại thư mục dự án và chạy lệnh: npm install\n"
        "Bước 5: Tạo file .env dựa trên file .env.example và điền các API Key (GEMINI_API_KEY, DATABASE_URL)."
    )

    # --- CHƯƠNG 9 ---
    doc.add_heading('CHƯƠNG 9: KHỞI TẠO DỮ LIỆU & DATABASE', level=1)
    doc.add_paragraph(
        "Sau khi cài đặt thư viện, thực hiện các lệnh sau để đồng bộ cơ sở dữ liệu:\n"
        "1. npx prisma generate: Khởi tạo Prisma Client.\n"
        "2. npx prisma db push: Đẩy cấu trúc bảng vào dữ liệu (hoặc npx prisma migrate dev).\n"
        "3. npm run db:seed: Chạy script nạp dữ liệu mẫu ban đầu (Môn học, Chương, Bài)."
    )

    # --- CHƯƠNG 10 ---
    doc.add_heading('CHƯƠNG 10: HƯỚNG DẪN SỬ DỤNG (USER MANUAL)', level=1)
    doc.add_paragraph(
        "Dành cho Giáo viên: Vào Dashboard -> Ngân hàng câu hỏi -> Tạo câu hỏi bằng AI. Chọn môn học và nhấn 'Generate'.\n"
        "Dành cho Học sinh: Vào Dashboard -> Đề thi. Chọn một bộ đề và bắt đầu làm bài. Sau khi hoàn thành, hãy xem bảng phân tích năng lực để biết mình cần ôn tập phần nào."
    )

    # --- CHƯƠNG 11 ---
    doc.add_heading('CHƯƠNG 11: BẢO MẬT & TỐI ƯU HÓA', level=1)
    doc.add_paragraph(
        "Hệ thống sử dụng NextAuth.js để quản lý phiên đăng nhập và bảo mật Role-based Access Control (RBAC). "
        "Mỗi người dùng chỉ có thể truy cập các tính năng phù hợp với quyền hạn của mình (Admin, Teacher, Student). "
        "Các API Routes được bảo mật bằng cơ chế kiểm tra Session tại phía Server."
    )

    # --- CHƯƠNG 12 ---
    doc.add_heading('CHƯƠNG 12: KẾT LUẬN & HƯỚNG PHÁT TRIỂN', level=1)
    doc.add_paragraph(
        "Dự án EduQuiz AI đã đạt được những kết quả ấn tượng trong việc ứng dụng AI vào giáo dục. "
        "Hướng phát triển tương lai bao gồm:\n"
        "- Xây dựng ứng dụng di động (Mobile App) hỗ trợ học tập mọi lúc mọi nơi.\n"
        "- Tích hợp AI tạo Slide bài giảng tự động từ ngân hàng câu hỏi.\n"
        "- Chế độ thi đấu PvP thời gian thực giữa các học sinh."
    )

    doc.add_paragraph("\n" * 2)
    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_p.add_run("--- HẾT ---")
    final_run.font.bold = True
    final_run.font.name = 'Times New Roman'

    filename = "EduQuiz_AI_Ultimate_Report.docx"
    doc.save(filename)
    print(f"Success: {filename} created.")

if __name__ == "__main__":
    create_ultimate_report()
